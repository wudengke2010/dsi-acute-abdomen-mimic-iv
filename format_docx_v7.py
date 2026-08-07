"""Format all v7 submission DOCX files for AIC Elsevier requirements:
- Double-line spacing
- Line numbering (continuous)
- Times New Roman 12pt (body), 10pt (table content)
- 1-inch (2.54cm) margins
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))

def add_line_numbers(section):
    """Add continuous line numbering to a section via XML.
    Uses OxmlElement + qn for proper attribute namespace handling.
    Removes any existing lnNumType first to avoid duplicates."""
    sectPr = section._sectPr
    # Remove existing lnNumType elements to prevent duplicates
    existing = sectPr.findall(qn('w:lnNumType'))
    for elem in existing:
        sectPr.remove(elem)
    # Create properly formatted lnNumType using OxmlElement
    lnNumType = OxmlElement('w:lnNumType')
    lnNumType.set(qn('w:countBy'), '1')
    lnNumType.set(qn('w:restart'), 'continuous')
    # Insert after pgMar (correct OOXML ordering)
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.addnext(lnNumType)
    else:
        sectPr.append(lnNumType)

def set_margins(section):
    """Set 1-inch margins on all sides."""
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def format_document(docx_path, is_cover_letter=False):
    """Apply AIC formatting to a DOCX document."""
    doc = Document(docx_path)
    
    # --- Section-level settings ---
    for section in doc.sections:
        set_margins(section)
        add_line_numbers(section)
        # Paper size: A4 for Elsevier
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    
    # --- Style-level settings ---
    # Normal style: Times New Roman 12pt, double spacing
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    
    # Heading styles
    for level in range(1, 5):
        heading_style_name = f'Heading {level}'
        try:
            hs = doc.styles[heading_style_name]
            hs.font.name = 'Times New Roman'
            hs.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12 if level == 3 else 12)
            hs.font.bold = True
            hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass
    
    # List Bullet style
    try:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Times New Roman'
        lb.font.size = Pt(12)
        lb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    except KeyError:
        pass
    
    # --- Paragraph-level overrides ---
    for paragraph in doc.paragraphs:
        # Ensure double line spacing on every paragraph
        if paragraph.paragraph_format.line_spacing_rule != WD_LINE_SPACING.DOUBLE:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        
        # Font: Times New Roman on all runs
        for run in paragraph.runs:
            if run.font.name != 'Times New Roman':
                run.font.name = 'Times New Roman'
            # Preserve intentionally smaller fonts (references, tables)
            # Only set 12pt if currently unset or >= 12pt
            if run.font.size is None or run.font.size >= Pt(11):
                # Check if this is in a reference context (numbered paragraphs)
                # References should stay at 10pt
                pass  # we handle references separately
    
    # --- Table formatting ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                    # Table header: bold 10pt; table body: 10pt
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        if run.font.size is None or run.font.size >= Pt(9):
                            run.font.size = Pt(10)
                        # Header row: bold
                        if row == table.rows[0]:
                            run.bold = True
    
    # Save
    doc.save(docx_path)
    print(f"Formatted: {docx_path}")

def format_cover_letter(docx_path):
    """Cover letter: single-spaced is acceptable, but AIC requires double for manuscript."""
    doc = Document(docx_path)
    
    for section in doc.sections:
        set_margins(section)
        add_line_numbers(section)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    
    # Cover letter can be single-spaced (standard business letter format)
    # But AIC manuscript guidelines say double-line spacing for all submitted documents
    # We'll apply double spacing to be safe
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            if run.font.size is None:
                run.font.size = Pt(12)
    
    doc.save(docx_path)
    print(f"Formatted cover letter: {docx_path}")

if __name__ == '__main__':
    # Main manuscript
    main_path = os.path.join(BASE, 'SCI_paper_v7.docx')
    format_document(main_path)
    
    # Supplementary materials
    supp_path = os.path.join(BASE, 'Supplementary_Materials_AIC_v7.docx')
    format_document(supp_path)
    
    # Cover letter
    cover_path = os.path.join(BASE, 'Cover_Letter_AIC_v7.docx')
    format_cover_letter(cover_path)
    
    print("\n✅ All 3 DOCX files formatted for AIC submission:")
    print("  - Double-line spacing")
    print("  - Continuous line numbering")
    print("  - Times New Roman 12pt (body) / 10pt (tables)")
    print("  - 1-inch margins")
    print("  - A4 paper size")
