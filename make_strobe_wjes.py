# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)

doc.add_heading('STROBE Statement - Checklist of Items for Cohort Studies', level=1)
p = doc.add_paragraph('')
r = p.add_run('Manuscript: Diastolic Shock Index as an Independent Predictor of In-Hospital Mortality in Critically Ill Patients with Acute Abdomen: A Retrospective Cohort Study with External Validation')
r.italic = True
doc.add_paragraph('Note: Location refers to named sections of the main manuscript (double-spaced, continuously line-numbered DOCX); tables and figures refer to those of the main text, with S-prefix denoting supplementary materials.')

items = [
    ("1a", "Title and abstract", "Indicate the study's design with a commonly used term in the title or the abstract", "Title: 'A Retrospective Cohort Study with External Validation'"),
    ("1b", "Title and abstract", "Provide in the abstract an informative and balanced summary of what was done and what was found", "Abstract (Background/Methods/Results/Conclusions, 349 words)"),
    ("2", "Introduction", "Explain the scientific background and rationale for the investigation being reported", "Background, paragraphs 1-3"),
    ("3", "Introduction", "State specific objectives, including any prespecified hypotheses", "Background, final paragraph (six prespecified aims)"),
    ("4", "Methods", "Present key elements of study design early in the paper", "Methods, Data sources and study design"),
    ("5", "Methods", "Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection", "Methods, Data sources and study design (MIMIC-IV v3.1, BIDMC Boston, 2008-2022; eICU-CRD v2.0, 208 US hospitals, 2014-2015)"),
    ("6a", "Methods", "Give eligibility criteria, and the sources and methods of selection of participants", "Methods, Study population (age >=18, ICU admission via ED, acute abdomen ICD-9/10 codes per Table S1, complete vital signs within 24h)"),
    ("6b", "Methods", "For matched studies, give matching criteria and number of exposed/unexposed", "Not applicable (unmatched cohort)"),
    ("7a", "Methods", "Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers", "Methods, Variables and statistical analysis (SI/MSI/DSI/Age-SI definitions; primary outcome in-hospital mortality; covariates age, sex, CCI, lactate, WBC, vasopressor, MV, SOFA)"),
    ("7b", "Methods", "Give criteria for classification of acute abdomen subtypes", "Methods, Study population (ICD-based: perforation > ischemia > obstruction > inflammation > other; Table S5)"),
    ("8", "Methods", "For each variable of interest, give sources of data and details of methods of assessment", "Methods, Variables and statistical analysis (first 24h ICU vital signs; BP priority: arterial line > NIBP > manual; mean 24h values)"),
    ("9", "Methods", "Describe any efforts to address potential sources of bias", "Methods (Study population; Variables and statistical analysis) and Discussion (Bias and causal inference): DAG-identified collider stratification bias; MI N=8,933; MICE; E-values; BP measurement hierarchy; survivorship bias of surgery covariate addressed by primary model without surgery"),
    ("10", "Methods", "Explain how the study size was arrived at", "Methods, Study population (all eligible admissions screened from 546,028; 5,728 complete cases; Figure 1)"),
    ("11", "Methods", "Explain how quantitative variables were handled in the analyses", "Methods, Variables and statistical analysis (mean 24h values; DSI quartiles and RCS with 4 knots for dose-response)"),
    ("12a", "Methods", "Describe all statistical methods, including those used to control for confounding", "Methods, Variables and statistical analysis (multivariable logistic regression at three nested model levels; ROC/AUC with DeLong tests; BH-FDR correction)"),
    ("12b", "Methods", "Describe any methods used to examine subgroups and interactions", "Methods, Variables and statistical analysis; Results, Sensitivity analyses (subtype-specific analyses, Figure S6; 14 sensitivity analyses, Table S11)"),
    ("12c", "Methods", "Explain how missing data were addressed", "Methods, Variables and statistical analysis (complete-case primary analysis; IterativeImputer MI on N=8,933; MICE with Rubin's rules as second strategy; Table S8 for excluded-patient comparison)"),
    ("12d", "Methods", "If applicable, explain how loss to follow-up was addressed", "Not applicable (in-hospital mortality ascertained via hospital_expire_flag; complete follow-up)"),
    ("12e", "Methods", "Describe any sensitivity analyses", "Methods, Variables and statistical analysis; Results, Sensitivity analyses; Table S11 (14 scenarios covering nine bias categories)"),
    ("13a", "Results", "Report numbers of individuals at each stage of the study", "Figure 1 (546,028 -> 8,933 eligible -> 5,728 complete cases)"),
    ("13b", "Results", "Give reasons for non-participation at each stage", "Methods, Study population; Figure 1; Table S8 (missing lactate/WBC; excluded patients less sick: mortality 8.0% vs 19.9%)"),
    ("14a", "Results", "Give characteristics of study participants and information on exposures and potential confounders", "Table 1 (baseline by DSI quartile); Table S7 (eICU baseline)"),
    ("14b", "Results", "Indicate number of participants with missing data for each variable of interest", "Methods, Study population; Table S8 (3,205 excluded, primarily missing lactate/WBC)"),
    ("15", "Results", "Report numbers of outcome events or summary measures", "Results, Baseline characteristics (1,141 in-hospital deaths, 19.9%; 758 ICU deaths, 13.2%)"),
    ("16a", "Results", "Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision", "Table 2 Panel A (adjusted ORs with 95% CIs); Table S6 (full coefficients); Figure 3 (forest plot)"),
    ("16b", "Results", "Report category boundaries when continuous variables were categorized", "Results, DSI quartile mortality gradient (DSI quartile cutoffs: Q1 <1.279, Q2 1.279-1.502, Q3 1.502-1.762, Q4 >1.762)"),
    ("16c", "Results", "If relevant, consider translating estimates of relative risk into absolute risk in meaningful ways", "Results, Primary model results (NRI, IDI, DCA net benefit, calibration, Brier score; Table S12 threshold diagnostic performance)"),
    ("17", "Results", "Report other analyses done - e.g., analyses of subgroups and interactions, sensitivity analyses", "Results, Component decomposition (Figure S8); Results, Sensitivity analyses and Table S11 (14 sensitivity analyses); Figure S6 (subtypes)"),
    ("18", "Discussion", "Summarize key results with reference to study objectives", "Discussion, opening paragraph (five principal findings)"),
    ("19", "Discussion", "Discuss limitations of the study, taking into account sources of potential bias or imprecision", "Discussion, Strengths and limitations (seven limitations addressed)"),
    ("20", "Discussion", "Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, results from similar studies, and other relevant evidence", "Discussion subsections: Bias and causal inference; Pathophysiological rationale; Comparison with previous studies; Clinical implications"),
    ("21", "Discussion", "Discuss the generalizability (external validity) of the study results", "Discussion, opening paragraph finding 5 (eICU-CRD validation across 208 hospitals; calibration requires local recalibration); Strengths and limitations"),
    ("22", "Other information", "Give the source of funding and the role of the funders for the present study and, if applicable, for the original study on which the present article is based", "Declarations, Funding (GWJJMB202510024181, kq2014242, 2021JJ30959; funders had no role)"),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, h in enumerate(['Item', 'Section (STROBE)', 'Recommendation', 'Location in Manuscript']):
    hdr[i].text = h
    for pgh in hdr[i].paragraphs:
        for run in pgh.runs:
            run.bold = True

for item in items:
    row = table.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]
    row[2].text = item[2]
    row[3].text = item[3]

doc.save('STROBE_Checklist_WJES.docx')
print('STROBE checklist saved:', len(items), 'items,', len(table.columns), 'columns')
