#!/usr/bin/env python3
"""Generate SCI_paper_v7.docx, Supplementary_Materials_AIC_v7.docx, and Cover_Letter_AIC_v7.docx using python-docx.

Based on v7 paper (SCI_paper_v7.md), Table_S*.csv files, and v6_revision_statistics.json.
Helper functions adapted from generate_docx_v6.py.
"""

import os
import json
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# Helper functions
# ============================================================

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_paragraph(doc, text, style='Normal', bold=False, italic=False,
                         font_size=12, font_name='Times New Roman',
                         alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         space_after=Pt(6), space_before=Pt(0),
                         first_line_indent=None, color=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East Asian font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_heading_styled(doc, text, level=1, font_size=14, font_name='Times New Roman'):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), font_name)
    return h


def add_table_from_data(doc, headers, rows, col_widths=None, header_color='4472C4',
                        font_size=10, bold_header=True, caption=None):
    """Add a formatted table with optional caption."""
    if caption:
        add_styled_paragraph(doc, caption, bold=True, font_size=11,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(4))

    n_cols = len(headers)
    n_rows = len(rows) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = bold_header
        run.font.size = Pt(font_size)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table


def enable_line_numbering(doc):
    """Enable line numbering in the document (continuous, restart per page)."""
    for section in doc.sections:
        sectPr = section._sectPr
        # Check if lnNumType already exists
        existing = sectPr.find(qn('w:lnNumType'))
        if existing is None:
            lnNumType = parse_xml(
                f'<w:lnNumType {nsdecls("w")} w:countBy="1" w:restart="newPage"/>'
            )
            sectPr.append(lnNumType)


def setup_document(doc, title=None, line_numbering=False):
    """Setup document defaults: Times New Roman, 12pt, double-spaced, 1-inch margins."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    if line_numbering:
        enable_line_numbering(doc)

    if title:
        add_styled_paragraph(doc, title, bold=True, font_size=16,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))


def read_csv_data(filename):
    """Read a CSV file and return headers + rows."""
    path = os.path.join(BASE_DIR, filename)
    with open(path, 'r', encoding='utf-8-sig') as f:
        reader = csv.reader(f)
        data = list(reader)
    return data[0], data[1:]


# ============================================================
# 1. Generate SCI_paper_v7.docx (Main Manuscript)
# ============================================================

def generate_main_paper():
    doc = Document()
    setup_document(doc, line_numbering=True)

    # Load statistics
    with open(os.path.join(BASE_DIR, 'v6_revision_statistics.json'), 'r') as f:
        v6_stats = json.load(f)
    with open(os.path.join(BASE_DIR, 'eicu_external_validation_results.json'), 'r') as f:
        eicu = json.load(f)

    # ===== TITLE PAGE =====
    add_styled_paragraph(doc,
        'Diastolic Shock Index as an Independent Predictor of In-Hospital Mortality '
        'in Critically Ill Patients with Acute Abdomen: A Retrospective Cohort Study '
        'with External Validation',
        bold=True, font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # Authors
    add_styled_paragraph(doc,
        'Jiqiang Liu [1]\u2020, Dengke Wu [1]*',
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    # Affiliations
    add_styled_paragraph(doc,
        '[1] Department of Emergency Medicine, and Emergency Medicine and Difficult Diseases Institute, '
        'The Second Xiangya Hospital of Central South University, Changsha 410011, Hunan, China',
        font_size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

    add_styled_paragraph(doc,
        '\u2020 First author.',
        font_size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

    add_styled_paragraph(doc,
        '* Corresponding author: Dengke Wu, Department of Emergency Medicine, and Emergency Medicine '
        'and Difficult Diseases Institute, The Second Xiangya Hospital of Central South University, '
        'Changsha 410011, Hunan, China. Electronic address: wudk2010@csu.edu.cn',
        font_size=10, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # ===== ABSTRACT =====
    add_heading_styled(doc, 'Abstract', level=1, font_size=14)

    # Background
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Background: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'Shock index (SI) derivatives predict mortality in trauma and sepsis but remain unexplored in '
        'acute abdomen. We evaluated SI-derived parameters for predicting in-hospital mortality, assessing '
        'diastolic shock index (DSI) as an independent, zero-cost bedside predictor complementary to SOFA.'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Methods
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Methods: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'Retrospective cohort from MIMIC-IV v3.1 (2008\u20132022). Adult ICU patients with acute abdomen '
        'ICD codes were included; SI/MSI/DSI/Age-SI were calculated from 24-hour vital signs. Primary '
        'outcome: in-hospital mortality. The primary model excluded surgery (survivorship bias), '
        'incorporating age, sex, Charlson Comorbidity Index (CCI), lactate, WBC, vasopressor use, '
        'mechanical ventilation, and SOFA. Performance was assessed via AUC/DeLong, NRI/IDI, DCA, RCS, '
        'cumulative incidence functions, calibration, and bootstrap validation. Multiple imputation '
        '(N=8,933) and 12 sensitivity analyses were performed. External validation used eICU-CRD v2.0 '
        '(N=5,755, 208 hospitals).'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Results
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Results: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'Among 5,728 complete-case ICU stays (median age 68, 56.0% male, in-hospital mortality 19.9%, '
        'SOFA 7 [4\u201311]), DSI was the strongest SI derivative. Extended baseline (no surgery) AUC=0.785; '
        'adding DSI yielded AUC=0.790 (\u0394AUC=0.005, DeLong P=0.012). DSI was an independent predictor '
        '(OR=2.18, 95% CI 1.79\u20132.65, P=7.59\u00d710\u207b\u00b9\u2075). \u0394AUC was below clinical relevance thresholds '
        '(\u22650.02), and categorical NRI crossed zero (0.008, 95% CI \u22120.009 to 0.044). However, category-free '
        'NRI (0.252, P<0.001) and IDI (0.013, P<0.001) were significant, and the DSI quartile gradient was '
        'dramatic: Q1=12.1%\u2192Q4=32.8% (2.7-fold, P=2.02\u00d710\u207b\u2074\u2079). Of 1,141 hospital deaths, 33.6% '
        'occurred after ICU discharge. External validation in eICU-CRD preserved discrimination '
        '(AUC=0.792, \u0394AUC=0.0074, DeLong P=0.0026) and replicated the quartile gradient '
        '(12.0%\u219233.5%), though calibration required local recalibration (intercept shift \u22123.935).'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Conclusions
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Conclusions: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'DSI is an independent predictor of in-hospital mortality after SOFA adjustment, providing '
        'zero-cost bedside risk stratification with a dramatic quartile gradient. While \u0394AUC is below '
        'clinical relevance thresholds and categorical NRI crosses zero, the independent OR, category-free '
        'NRI, and quartile gradient support DSI as a complementary, immediately available tool when '
        'laboratory data are pending. Prediction was most pronounced in non-surgical acute abdomen '
        '(AUC=0.826).'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Keywords
    add_styled_paragraph(doc,
        'Keywords: Diastolic shock index; Acute abdomen; In-hospital mortality; SOFA; External validation',
        italic=True, font_size=12, space_after=Pt(12))

    # ===== ABBREVIATIONS =====
    add_heading_styled(doc, 'Abbreviations', level=1, font_size=14)
    add_styled_paragraph(doc,
        'SI, shock index; MSI, modified shock index; DSI, diastolic shock index; Age-SI, age-adjusted '
        'shock index; HR, heart rate; SBP, systolic blood pressure; DBP, diastolic blood pressure; MAP, '
        'mean arterial pressure; SOFA, Sequential Organ Failure Assessment; CCI, Charlson Comorbidity '
        'Index; ICU, intensive care unit; LOS, length of stay; ROC, receiver operating characteristic; '
        'AUC, area under the curve; NRI, net reclassification improvement; IDI, integrated discrimination '
        'improvement; cf-NRI, category-free NRI; cat-NRI, categorical NRI; DCA, decision curve analysis; '
        'RCS, restricted cubic splines; CIF, cumulative incidence function; HL, Hosmer-Lemeshow; VIF, '
        'variance inflation factor; MV, mechanical ventilation; ED, emergency department; ICD, '
        'International Classification of Diseases; WBC, white blood cell count; MI, multiple imputation',
        font_size=11, space_after=Pt(12))

    # ===== 1. INTRODUCTION =====
    add_heading_styled(doc, '1. Introduction', level=1, font_size=14)

    add_styled_paragraph(doc,
        'Acute abdomen\u2014severe abdominal pain of sudden onset requiring urgent evaluation\u2014remains '
        'one of the most challenging presentations in emergency medicine [1]. Early risk stratification '
        'is critical, as outcomes vary dramatically across etiologies [1,2].',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'The shock index (SI=HR/SBP), first described by Allg\u00f6wer and Burri [3], has inspired '
        'derivatives: modified shock index (MSI=HR/MAP) [4], diastolic shock index (DSI=HR/DBP) [5], '
        'and age-adjusted shock index (Age-SI=SI\u00d7Age/10) [6]. These predict mortality in trauma [7] '
        'and sepsis [8], but have never been systematically evaluated in acute abdomen\u2014a population '
        'with pathophysiological diversity including inflammation-driven vasodilation, mechanical '
        'obstruction, perforation-induced peritonitis, and ischemia.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Previous SI-derivative studies relied on AUC comparisons without evaluating independent '
        'predictive value beyond established ICU predictors, nor assessing robustness via bootstrap '
        'validation, sensitivity analyses, or competing risks. The TRIPOD+AI guidelines [9] emphasize '
        'that prediction models must demonstrate clinical benefit via NRI/IDI/DCA [10] and undergo '
        'external validation. This study was reported per STROBE [11] and TRIPOD+AI guidelines [9].',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'We aimed to: (1) compare SI, MSI, DSI, and Age-SI for in-hospital mortality prediction; '
        '(2) evaluate DSI as an independent predictor beyond SOFA-adjusted covariates; (3) assess '
        'robustness via bootstrap, MI, and 12 sensitivity analyses; (4) externally validate in '
        'eICU-CRD; and (5) determine subtype-specific performance.',
        font_size=12, first_line_indent=Inches(0.5))

    # ===== 2. METHODS =====
    add_heading_styled(doc, '2. Methods', level=1, font_size=14)

    # 2.1
    add_heading_styled(doc, '2.1 Data Sources and Study Design', level=2, font_size=12)
    add_styled_paragraph(doc,
        'This retrospective cohort study utilized MIMIC-IV v3.1 [12] (Beth Israel Deaconess Medical '
        'Center, Boston, 2008\u20132022; 546,028 admissions). External validation used eICU-CRD v2.0 [13] '
        '(208 US hospitals, 2014\u20132015). Both accessed via PhysioNet with required training. IRB '
        'approval (BIDMC, MIT) was obtained; individual consent waived for de-identified data.',
        font_size=12, first_line_indent=Inches(0.5))

    # 2.2
    add_heading_styled(doc, '2.2 Study Population (Figure 1)', level=2, font_size=12)

    add_styled_paragraph(doc,
        'Inclusion: age \u226518; ICU admission via ED; acute abdomen ICD-9/10 codes (Supplementary '
        'Table S1); complete vital signs (HR, SBP, DBP) within 24h. Exclusion: missing extended '
        'covariates (lactate, WBC). From 546,028 admissions, 5,728 complete cases (CC) were analyzed. '
        'The 3,205 excluded patients had substantially lower severity (mortality 8.0% vs 19.9%, '
        'vasopressor 12% vs 43.6%, MV 17% vs 52.5%), reflecting selection bias toward patients '
        'receiving arterial blood gas monitoring (Supplementary Table S8). MI (N=8,933) addressed '
        'this bias.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Acute abdomen subtypes based on ICD: perforation, ischemia, obstruction, inflammation, and '
        'other (29.9%, N=1,712, containing complications alongside primary codes; Supplementary '
        'Table S5). Priority: perforation > ischemia > obstruction > inflammation > other.',
        font_size=12, first_line_indent=Inches(0.5))

    # 2.3
    add_heading_styled(doc, '2.3 Variables and Analysis', level=2, font_size=12)

    add_styled_paragraph(doc,
        'SI derivatives (24h vital signs): SI=HR/SBP, MSI=HR/MAP, DSI=HR/DBP, Age-SI=SI\u00d7Age/10. '
        'BP priority: arterial > non-invasive > manual. Primary outcome: in-hospital mortality '
        '(hospital_expire_flag); secondary: ICU mortality.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Primary extended baseline (without surgery): age, sex, CCI [14], lactate, WBC, vasopressor '
        'use, MV, SOFA [15]. Surgery was excluded because \u201csurgery during hospitalization\u201d (67.4%) '
        'introduces survivorship bias: only 5.1% had surgery \u226424h from ICU admission, confirming most '
        'procedures occurred after surviving the acute crisis (surgery_24h OR=0.88, P=0.46). A model '
        'including surgery is reported as an alternative (Supplementary Table S9). A parsimonious model '
        '(without vasopressor and MV) is also reported (Supplementary Table S10), since these covariates '
        'were non-significant after SOFA adjustment (P=0.14 and 0.45 respectively).',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Statistical methods: Multivariable logistic regression at three levels (basic, extended, '
        'extended+DSI). ROC/AUC with DeLong comparisons [16]. VIF for multicollinearity (all <3.0; '
        'maximum SOFA=2.42). NRI: categorical (<10%, 10\u201330%, >30% risk strata) and category-free [10]; '
        'IDI [10]. DCA [17]. RCS (4 knots) [18]. Cumulative incidence functions for competing risks [19]; '
        'formal Fine-Gray models were not implemented (CIF curves are descriptive only). Calibration: HL '
        'test, Brier score. Bootstrap validation (200 resamples; optimism=0.002). Sensitivity analyses '
        '(12 scenarios, Supplementary Table S11). MI: 5 imputations on N=8,933.',
        font_size=12, first_line_indent=Inches(0.5))

    # 2.4
    add_heading_styled(doc, '2.4 External Validation', level=2, font_size=12)
    add_styled_paragraph(doc,
        'MIMIC-IV models were applied to eICU-CRD without retraining (TRIPOD type 2b/3b). '
        'Methodological note: eICU SOFA was computed from APACHE APS variables rather than identical '
        'MIMIC-IV concept definitions; platelets were unavailable (hematocrit substituted). This '
        'heterogeneity is reflected in higher eICU median SOFA (9 [7\u201312] vs MIMIC-IV 7 [4\u201311]). '
        'Both un-recalibrated and recalibrated (logistic intercept/slope adjustment) metrics are '
        'reported per TRIPOD+AI [9]. Performance: AUC/DeLong, cf-NRI/IDI, DSI quartile gradient '
        '(derivation cutoffs applied; eICU quartile sizes unequal: Q1=1,294, Q4=1,677). All analyses: '
        'Python 3.13 (pandas, scipy, statsmodels, DuckDB).',
        font_size=12, first_line_indent=Inches(0.5))

    # ===== 3. RESULTS =====
    add_heading_styled(doc, '3. Results', level=1, font_size=14)

    # 3.1
    add_heading_styled(doc, '3.1 Baseline Characteristics (Table 1)', level=2, font_size=12)
    add_styled_paragraph(doc,
        'Among 5,728 CC ICU stays: median age 68 [57\u201379], 56.0% male, in-hospital mortality 19.9% '
        '(1,141 deaths; 33.6% after ICU discharge). SOFA 7 [4\u201311]; significantly higher in non-survivors '
        '(11 [8\u201315] vs 6 [4\u201310]). Vasopressor use 43.6%; MV 52.5%; lactate 2.0 [1.3\u20133.2] mmol/L; '
        'CCI 3 [1\u20135]. Subtypes: inflammation (37.5%), other (29.9%), obstruction (20.6%), ischemia '
        '(6.2%), perforation (5.8%).',
        font_size=12, first_line_indent=Inches(0.5))

    # Table 1
    add_table_from_data(doc,
        headers=['Characteristic', 'Overall (N=5,728)', 'Q1 (Low, n=1,432)',
                 'Q2 (n=1,432)', 'Q3 (n=1,432)', 'Q4 (High, n=1,432)'],
        rows=[
            ['Age, years', '68.0 [57.0\u201379.0]', '66.0 [55.0\u201378.0]', '68.0 [57.0\u201379.0]',
             '69.0 [58.0\u201380.0]', '69.0 [56.0\u201380.0]'],
            ['Male, n (%)', '3,206 (56.0%)', '897 (62.6%)', '800 (55.9%)', '796 (55.6%)', '713 (49.8%)'],
            ['SOFA score', '7.0 [4.0\u201311.0]', '6.0 [3.0\u20139.0]', '7.0 [4.0\u201310.0]',
             '8.0 [5.0\u201311.0]', '10.0 [6.0\u201313.0]'],
            ['Lactate, mmol/L', '2.0 [1.3\u20133.2]', '1.7 [1.2\u20132.7]', '1.9 [1.2\u20132.9]',
             '2.0 [1.4\u20133.3]', '2.6 [1.5\u20134.1]'],
            ['Vasopressor use, n (%)', '2,498 (43.6%)', '395 (27.6%)', '562 (39.2%)',
             '678 (47.3%)', '863 (60.3%)'],
            ['Mechanical ventilation, n (%)', '3,007 (52.5%)', '601 (42.0%)', '705 (49.2%)',
             '799 (55.8%)', '902 (63.0%)'],
            ['In-hospital mortality, n (%)', '1,141 (19.9%)', '173 (12.1%)', '208 (14.5%)',
             '291 (20.3%)', '469 (32.8%)'],
        ],
        caption='Table 1. Baseline characteristics by DSI quartile',
        font_size=9)

    # 3.2
    add_heading_styled(doc, '3.2 DSI Quartile Mortality Gradient (Table 1)', level=2, font_size=12)
    add_styled_paragraph(doc,
        'DSI (mean 24h) quartiles demonstrated a dramatic gradient (\u03c7\u00b2=229, P=2.02\u00d710\u207b\u2074\u2079): '
        'Q1<1.279 (12.1%), Q2 1.279\u20131.502 (14.5%), Q3 1.502\u20131.762 (20.3%), Q4>1.762 (32.8%)\u2014'
        'a 2.7-fold increase. Higher quartiles had progressively higher lactate, vasopressor use, and '
        'MV rates.',
        font_size=12, first_line_indent=Inches(0.5), space_before=Pt(6))

    # 3.3
    add_heading_styled(doc, '3.3 Primary Model Results (Table 2)', level=2, font_size=12)

    add_styled_paragraph(doc,
        'Table 2, Panel A (MIMIC-IV): Extended baseline (no surgery) AUC=0.785 [0.769\u20130.801]; '
        'adding DSI: AUC=0.790 [0.775\u20130.805], \u0394AUC=0.005 (DeLong P=0.012). DSI: OR=2.18 [1.79\u20132.65], '
        'P=7.59\u00d710\u207b\u00b9\u2075; SOFA: OR=1.16 [1.13\u20131.19], P<10\u207b\u00b3\u2076; lactate: OR=1.14 [1.11\u20131.17]; '
        'CCI: OR=1.14 [1.11\u20131.16]. Vasopressor (P=0.14), MV (P=0.45), gender (P=0.08), and WBC '
        '(P=0.07) were non-significant after SOFA adjustment. \u0394AUC=0.005 was below clinical relevance '
        'thresholds (\u22650.02 per Cook [20] and Vickers [21]); categorical NRI (0.008, CI \u22120.009 to '
        '0.044) crossed zero. Category-free NRI (0.252, CI 0.183\u20130.331) and IDI (0.013, CI '
        '0.007\u20130.020) were significant, confirming additional continuous prognostic information. '
        'VIF all <3.0; bootstrap optimism=0.002. Full coefficients: Supplementary Table S6; Forest '
        'plot: Figure 3; ROC curves: Figure 2.',
        font_size=12, first_line_indent=Inches(0.5))

    # Table 2
    add_table_from_data(doc,
        headers=['Model', 'AUC (95% CI)', '\u0394AUC', 'DeLong P', 'cat-NRI (95% CI)',
                 'cf-NRI (95% CI)', 'IDI (95% CI)'],
        rows=[
            ['Extended baseline (no surgery)', '0.785 (0.769\u20130.801)', '\u2014', '\u2014', '\u2014', '\u2014', '\u2014'],
            ['Extended + DSI (primary)', '0.790 (0.775\u20130.805)', '0.005', '0.012',
             '0.008 (\u22120.009, 0.044)', '0.252 (0.183, 0.331)', '0.013 (0.007, 0.020)'],
        ],
        caption='Table 2, Panel A. Primary model performance (MIMIC-IV, without surgery)',
        font_size=9)

    add_styled_paragraph(doc,
        'Table 2, Panel B (eICU-CRD): N=5,755 (208 hospitals), mortality 20.0%, SOFA 9 [7\u201312]. '
        'Extended baseline AUC=0.785; Extended+DSI AUC=0.792; \u0394AUC=0.0074 (DeLong P=0.0026). '
        'Un-recalibrated: Brier=0.38\u20130.59, HL P<0.001 (catastrophically poor). After logistic '
        'recalibration (intercept \u22123.935, slope 0.952): Brier=0.126, HL P=0.266. cf-NRI=0.277, '
        'IDI=0.014. DSI quartile gradient: Q1=12.0%\u2192Q4=33.5% (2.8-fold). The large intercept shift '
        'indicates discrimination is transportable (slope near 1.0), but absolute risk estimates '
        'require local recalibration.',
        font_size=12, first_line_indent=Inches(0.5), space_before=Pt(6))

    add_table_from_data(doc,
        headers=['Metric', 'MIMIC-IV (Derivation)', 'eICU-CRD (Validation)'],
        rows=[
            ['N (CC)', '5,728', '5,755'],
            ['In-hospital mortality', '19.9%', '20.0%'],
            ['Median SOFA [IQR]', '7 [4\u201311]', '9 [7\u201312]'],
            ['Extended baseline AUC', '0.785', '0.785'],
            ['Extended+DSI AUC', '0.790', '0.792'],
            ['\u0394AUC', '0.005', '0.0074'],
            ['DeLong P', '0.012', '0.0026'],
            ['Brier (recalibrated)', '0.127', '0.126'],
            ['HL P (recalibrated)', '0.691', '0.266'],
            ['Recal intercept shift', '\u2014', '\u22123.935'],
            ['Recal slope', '\u2014', '0.952'],
            ['cf-NRI', '0.252', '0.277'],
            ['IDI', '0.013', '0.014'],
            ['DSI Q1 mortality', '12.1%', '12.0%'],
            ['DSI Q4 mortality', '32.8%', '33.5%'],
        ],
        caption='Table 2, Panel B. External validation: eICU-CRD vs MIMIC-IV',
        font_size=9)

    # 3.4
    add_heading_styled(doc, '3.4 Sensitivity Analyses (Supplementary Table S11)', level=2, font_size=12)
    add_styled_paragraph(doc,
        'DSI\'s independent predictive value was robust across 12 analyses (OR range 2.15\u20132.65, '
        'all P<10\u207b\u00b9\u00b9). Key findings: (1) Including surgery increased AUC by only 0.002 (baseline '
        '0.785\u21920.787); surgery appeared protective (OR=0.68) but likely reflects survivorship bias; '
        'surgery_24h (OR=0.88, P=0.46) was non-significant. (2) Excluding \u201cOther\u201d subtype (N=4,016): '
        'AUC=0.786\u20130.788, DSI OR=2.15\u20132.22. (3) Parsimonious model (without vasopressor/MV): '
        'AUC=0.789, DSI OR=2.22, \u0394AUC=0.005\u2014confirming redundant covariates do not affect DSI\'s '
        'value. (4) Non-surgical subgroup: AUC=0.826 (best performance). (5) MI (N=8,933): AUC=0.822, '
        'DSI OR=2.65, addressing selection bias with higher estimates. (6) RCS: all P_overall<0.001, '
        'P_nonlinear>0.05 (Supplementary Figure S3). (7) DCA: modest incremental net benefit '
        '(Supplementary Figure S4). (8) Calibration: Extended+DSI HL P=0.691, Brier=0.126 '
        '(Supplementary Figure S5). (9) Subtype-specific AUC: inflammation 0.819, obstruction 0.749, '
        'perforation 0.766, ischemia 0.807 (Supplementary Figure S6). (10) CIF curves confirmed '
        'progressive risk across quartiles (Supplementary Figure S7).',
        font_size=12, first_line_indent=Inches(0.5), space_before=Pt(6))

    # ===== 4. DISCUSSION =====
    add_heading_styled(doc, '4. Discussion', level=1, font_size=14)

    add_styled_paragraph(doc,
        'This study provides the first comprehensive evaluation of SI-derived parameters in acute '
        'abdomen ICU patients, with SOFA adjustment, bootstrap validation, MI, 12 sensitivity '
        'analyses, competing risk framework, external validation in 208 hospitals, and '
        'STROBE/TRIPOD+AI-compliant reporting.',
        font_size=12, first_line_indent=Inches(0.5))

    # Five findings
    add_styled_paragraph(doc,
        'First, DSI is an independent predictor after SOFA adjustment (OR=2.18, P=7.59\u00d710\u207b\u00b9\u2075). '
        'However, \u0394AUC=0.005 is below clinical relevance thresholds (\u22650.02) [20,21], and categorical '
        'NRI crossing zero means DSI does not reclassify patients across the 10%/30% risk strata beyond '
        'a model already containing SOFA and lactate. This is expected when a marker refines continuous '
        'prediction without shifting categorical thresholds [21]. We position DSI not as a SOFA '
        'replacement, but as a complementary zero-cost bedside tool providing independent risk '
        'information from routinely monitored HR and DBP\u2014available without laboratory turnaround time. '
        'DSI\'s clinical value lies in immediate risk stratification when SOFA components (platelets, '
        'bilirubin, PaO\u2082, vasopressor doses) are pending.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Second, the DSI quartile gradient (12.1%\u219232.8%, 2.7-fold) provides clinically actionable '
        'thresholds (Q1<1.279, Q4>1.762). Higher quartiles had progressively higher lactate, '
        'vasopressor, and MV rates, confirming DSI as an integrative hemodynamic severity marker.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Third, 33.6% of hospital deaths occurred after ICU discharge, justifying in-hospital mortality '
        'as the primary endpoint and suggesting DSI could guide post-ICU monitoring intensity.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Fourth, surgery was excluded from the primary model due to survivorship bias: 67.4% \u201cduring '
        'hospitalization\u201d vs 5.1% \u226424h, with surgery_24h non-significant (P=0.46). Including surgery '
        'increased baseline AUC by only 0.002, confirming its minimal and biased contribution. This bias '
        'affects any model including \u201csurgery during hospitalization\u201d as a covariate\u2014a caution for '
        'future ICU prediction studies.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Fifth, external validation in eICU-CRD (208 hospitals) preserved discrimination (AUC=0.792, '
        '\u0394AUC=0.0074) and replicated the quartile gradient (12.0%\u219233.5%). However, direct application '
        'yielded catastrophically poor calibration (Brier 0.38\u20130.59), requiring logistic recalibration '
        '(intercept \u22123.935, slope 0.952). The near-ideal slope confirms discrimination transportability, '
        'but the large intercept shift means MIMIC-IV-derived absolute risk estimates cannot be directly '
        'applied to new settings without local recalibration. Additionally, eICU SOFA computation differed '
        'from MIMIC-IV (APACHE APS-based, median 9 vs 7; hematocrit substituted for platelets), and eICU '
        'data (2014\u20132015) represents an older practice era.',
        font_size=12, first_line_indent=Inches(0.5))

    # Pathophysiological rationale
    add_heading_styled(doc, 'Pathophysiological Rationale', level=2, font_size=12)
    add_styled_paragraph(doc,
        'DSI captures HR-to-diastolic pressure relationships. In acute abdomen, progressive vasodilation '
        'and splanchnic compromise first manifest as diastolic pressure decline\u2014loss of peripheral '
        'vascular tone before systolic compensatory mechanisms fail\u2014making DSI more sensitive to early '
        'deterioration than SI (HR/SBP) [5].',
        font_size=12, first_line_indent=Inches(0.5))

    # Clinical implications
    add_heading_styled(doc, 'Clinical Implications', level=2, font_size=12)
    add_styled_paragraph(doc,
        'DSI provides: (1) independent predictive value (OR=2.18); (2) dramatic quartile gradient; '
        '(3) zero-cost from routine vitals; (4) validated discrimination across 208 hospitals. It '
        'complements but does not replace SOFA-based categorical risk classifications. In ischemia '
        '(mortality 40.5%, AUC=0.807) and non-surgical acute abdomen (AUC=0.826), DSI may be particularly '
        'useful.',
        font_size=12, first_line_indent=Inches(0.5))

    # Limitations
    add_heading_styled(doc, '4.1 Limitations', level=2, font_size=12)
    limitations = [
        '(1) Single-center retrospective derivation, though externally validated in 208 hospitals;',
        '(2) Selection bias: 36% exclusion (primarily lactate) enriched CC with more severe patients '
        '(mortality 19.9% vs 8.0% in excluded); MI (N=8,933) addressed this; eICU validation also used '
        'CC (32.6% rate);',
        '(3) \u0394AUC below clinical thresholds; categorical NRI crossing zero; DSI is positioned as a '
        'complementary bedside tool, not a SOFA replacement;',
        '(4) Surgery survivorship bias addressed by removal from primary model;',
        '(5) \u201cOther\u201d subtype heterogeneity (29.9%); sensitivity excluding it preserved results;',
        '(6) No Fine-Gray subdistribution hazard models; CIF curves are descriptive;',
        '(7) eICU SOFA heterogeneity (APS-based vs MIMIC-IV concept; median 9 vs 7); hematocrit '
        'substituted for platelets;',
        '(8) Large recalibration intercept shift (\u22123.935) means absolute risk predictions require '
        'local recalibration; discrimination (slope 0.952) is transportable;',
        '(9) eICU data (2014\u20132015) represents older practice era;',
        '(10) WBC borderline (P=0.07); vasopressor and MV non-significant after SOFA adjustment '
        '(absorbed by SOFA components);',
        '(11) Only two authors; statistical expertise was guided by TRIPOD+AI guidelines and established '
        'biostatistical references [9,16,20,21].',
    ]
    for lim in limitations:
        add_styled_paragraph(doc, lim, font_size=12, first_line_indent=Inches(0.5))

    # Future Directions
    add_heading_styled(doc, '4.2 Future Directions', level=2, font_size=12)
    add_styled_paragraph(doc,
        'Prospective multicenter validation with standardized SOFA; Fine-Gray modeling; DSI trajectory '
        'analysis; integration with machine learning; non-US population validation.',
        font_size=12, first_line_indent=Inches(0.5))

    # ===== 5. CONCLUSIONS =====
    add_heading_styled(doc, '5. Conclusions', level=1, font_size=14)
    add_styled_paragraph(doc,
        'DSI is an independent predictor of in-hospital mortality in acute abdomen after SOFA adjustment '
        '(OR=2.18, 95% CI 1.79\u20132.65), providing zero-cost bedside risk stratification with a dramatic '
        'quartile gradient (2.7-fold). While \u0394AUC=0.005 is below clinical relevance thresholds and '
        'categorical NRI crosses zero, the independent OR, category-free NRI, and quartile gradient '
        'support DSI as a complementary, immediately available risk-stratification tool when laboratory '
        'data are pending. External validation in eICU-CRD (208 hospitals) confirmed discrimination '
        'transportability and replicated the quartile gradient, though calibration required local '
        'recalibration. Prediction was most pronounced in non-surgical acute abdomen (AUC=0.826).',
        font_size=12, first_line_indent=Inches(0.5))

    # ===== SUPPLEMENTARY MATERIALS REFERENCE =====
    add_heading_styled(doc, 'Supplementary Materials', level=1, font_size=14)
    supp_items = [
        'Table S1: ICD-9/10 codes for acute abdomen identification and subtype classification.',
        'Table S2: STROBE checklist (completed).',
        'Table S3: Baseline characteristics (N=5,728) by DSI quartile.',
        'Table S4: TRIPOD+AI checklist (27 items) [9].',
        'Table S5: ICD code composition of \u201cOther\u201d subtype (N=1,712).',
        'Table S6: Full model coefficients (primary model without surgery + DSI).',
        'Table S7: eICU-CRD baseline characteristics (N=5,755) by DSI quartile.',
        'Table S8: Excluded (N=3,205) vs CC (N=5,728) patient characteristics.',
        'Table S9: Alternative model (with surgery) coefficients and performance.',
        'Table S10: Parsimonious model (without vasopressor/MV) performance.',
        'Table S11: Sensitivity analyses summary (12 scenarios).',
        'Figure S1: Calibration plots for basic baseline models.',
        'Figure S2: Kaplan-Meier curves by DSI quartile.',
        'Figure S3: RCS dose-response curves for SI, MSI, DSI, Age-SI.',
        'Figure S4: DCA net benefit across threshold probabilities.',
        'Figure S5: Calibration plots (4 model levels).',
        'Figure S6: Subgroup ROC curves by subtype.',
        'Figure S7: Cumulative incidence functions by DSI quartile.',
    ]
    for item in supp_items:
        add_styled_paragraph(doc, item, font_size=11, space_after=Pt(2))

    # ===== FIGURE LEGENDS =====
    add_heading_styled(doc, 'Figure Legends', level=1, font_size=14)
    fig_legends = [
        'Figure 1: Study flow diagram (546,028\u21925,728 CC), with excluded patient characteristics '
        'in Supplementary Table S8.',
        'Figure 2: ROC curves\u2014basic baseline, extended baseline (no surgery), extended+DSI (no '
        'surgery), extended+all SI derivatives.',
        'Figure 3: Forest plot of adjusted ORs from the primary model (extended baseline without '
        'surgery + DSI).',
    ]
    for legend in fig_legends:
        add_styled_paragraph(doc, legend, font_size=12, space_after=Pt(4))

    # ===== DECLARATIONS =====
    add_heading_styled(doc, 'Declarations', level=1, font_size=14)

    # Ethics
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Ethics: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'MIMIC-IV and eICU-CRD are publicly available with IRB approval (BIDMC, MIT). '
        'Individual consent waived for de-identified data.'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Consent for publication
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Consent for publication: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run('Not applicable.')
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Funding
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Funding: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'GWJJMB202510024181 (National Health Commission), kq2014242 (Changsha Science and Technology '
        'Bureau), 2021JJ30959 (Hunan Provincial Natural Science Foundation). Funders had no role in '
        'study design, analysis, or publication.'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Conflicts
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Conflicts: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run('Authors declare no conflicts.')
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # CRediT
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('CRediT: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'Jiqiang Liu: Conceptualization, Data curation, Formal analysis, Investigation, Methodology, '
        'Software, Validation, Visualization, Writing \u2013 original draft. Dengke Wu: Conceptualization, '
        'Funding acquisition, Methodology, Project administration, Resources, Supervision, '
        'Writing \u2013 review & editing.'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # AI use declaration
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('AI use declaration: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'During the preparation of this manuscript, the authors used a large language model for '
        'language polishing and manuscript editing only. All data extraction, statistical analyses, '
        'figure generation, and scientific interpretation were performed independently by the authors. '
        'After using the tool, the authors reviewed and edited the content as needed and take full '
        'responsibility for the content of the publication.'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Acknowledgments
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Acknowledgments: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'We thank the MIMIC-IV and eICU-CRD teams for open access to clinical databases.'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # Data availability
    p = add_styled_paragraph(doc, '', font_size=12, space_after=Pt(6))
    run = p.add_run('Data availability: ')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        'MIMIC-IV v3.1 at https://physionet.org/content/mimiciv/3.1/. eICU-CRD v2.0 at '
        'https://physionet.org/content/eicu-crd/2.0/. Analysis code available on request from '
        'the corresponding author.'
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # ===== REFERENCES =====
    add_heading_styled(doc, 'References', level=1, font_size=14)

    refs = [
        '1. Flum DR. Acute abdomen. In: Sabiston Textbook of Surgery. 21st ed. Elsevier; 2022.',
        '2. Cervero F, Laird JM. Visceral pain. Lancet. 1999;353(9170):2145-2148.',
        '3. Allg\u00f6wer M, Burri C. Schockindex. Deutsche Med Wochenschr. 1967;92(43):1947-1950.',
        '4. Jouffroy R, Gille S, Gilbert B, et al. Shock index derivatives and 28-day mortality in '
        'prehospital septic shock. J Emerg Med. 2024;66(2):144-153.',
        '5. Ospina-Tasc\u00f3n GA, Teboul JL, Hernandez G, et al. Diastolic shock index and clinical '
        'outcomes in septic shock. Ann Intensive Care. 2020;10:41.',
        '6. Kim SY, Hong KJ, Shin SD, et al. Validation of shock indices for predicting geriatric '
        'trauma mortality. J Korean Med Sci. 2016;31(12):2026-2032.',
        '7. Olaussen A, Peterson G, Synnot A, et al. Shock index and mortality in trauma: systematic '
        'review. Crit Care. 2023;27:88.',
        '8. Liu YC, Lee CT, Su HY, et al. Shock indices and in-hospital mortality in sepsis. '
        'PLoS One. 2024;19(3):e0298617.',
        '9. Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement. BMJ. 2024;385:e078378.',
        '10. Pencina MJ, D\'Agostino RB, et al. Evaluating added predictive ability. Stat Med. '
        '2008;27(2):157-172.',
        '11. von Elm E, Altman DG, Egger M, et al. STROBE statement. Lancet. '
        '2007;370(9596):1453-1457.',
        '12. Johnson AEW, Bulgarelli L, Pollard TJ, et al. MIMIC-IV. Sci Data. 2023;10:1.',
        '13. Pollard TJ, Johnson AEW, Raffa JD, et al. eICU-CRD. Sci Data. 2018;5:180175.',
        '14. Charlson ME, Pompei P, Ales KL, MacKenzie CR. Comorbidity classification. '
        'J Chronic Dis. 1987;40(5):373-383.',
        '15. Vincent JL, Moreno R, Takala J, et al. SOFA score. Intensive Care Med. '
        '1996;22(7):707-710.',
        '16. DeLong ER, DeLong DM, Clarke-Pearson DL. Comparing AUCs. Biometrics. '
        '1988;44(3):837-845.',
        '17. Vickers AJ, Elkin EB. Decision curve analysis. Med Decis Making. '
        '2006;26(6):565-574.',
        '18. Desquilbet L, Mariotti F. Dose-response via RCS. Am J Epidemiol. '
        '2010;172(12):1377-1385.',
        '19. Fine JP, Gray RJ. Proportional hazards model for competing risks. '
        'J Am Stat Assoc. 1999;94(446):496-509.',
        '20. Cook NR. Use and misuse of the receiver operating characteristic curve in risk prediction. '
        'Circulation. 2007;115(7):928-935.',
        '21. Vickers AJ, Cronin AM, Begg CB. One statistical test is sufficient for assessing '
        'prediction model performance. Med Decis Making. 2008;28(5):525-529.',
    ]
    for ref in refs:
        add_styled_paragraph(doc, ref, font_size=10, space_after=Pt(2))

    # Save
    out_path = os.path.join(BASE_DIR, 'SCI_paper_v7.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


# ============================================================
# 2. Generate Supplementary_Materials_AIC_v7.docx
# ============================================================

def generate_supplementary():
    doc = Document()
    setup_document(doc)

    add_styled_paragraph(doc,
        'Supplementary Materials',
        bold=True, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_styled_paragraph(doc,
        'Diastolic Shock Index as an Independent Predictor of In-Hospital Mortality '
        'in Critically Ill Patients with Acute Abdomen: A Retrospective Cohort Study '
        'with External Validation',
        bold=True, font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_styled_paragraph(doc,
        'Jiqiang Liu, Dengke Wu',
        font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))

    # Load statistics for Tables S6, S9, S10, S11
    with open(os.path.join(BASE_DIR, 'v6_revision_statistics.json'), 'r') as f:
        v6_stats = json.load(f)

    var_names = {
        'age_at_admission': 'Age (per year)',
        'gender_binary': 'Male gender',
        'CCI': 'CCI (per point)',
        'lactate_first': 'Lactate (per mmol/L)',
        'wbc_first': 'WBC (per \u00d710\u2079/L)',
        'vasopressor_use': 'Vasopressor use',
        'any_surgery': 'Surgery (during hospitalization)',
        'surgery_24h': 'Surgery (\u226424h)',
        'mechanical_ventilation': 'Mechanical ventilation',
        'sofa': 'SOFA (per point)',
        'DSI_mean': 'DSI (mean 24h)',
    }

    # ===== Table S1 =====
    add_heading_styled(doc,
        'Table S1. ICD Codes for Acute Abdomen Identification and Subtype Classification',
        level=2, font_size=11)
    headers, rows = read_csv_data('Table_S1_ICD_Codes.csv')
    add_table_from_data(doc, headers=headers, rows=rows, font_size=8)

    # ===== Table S2 =====
    add_heading_styled(doc, 'Table S2. STROBE Checklist', level=2, font_size=11)
    headers, rows = read_csv_data('Table_S2_STROBE_Checklist.csv')
    add_table_from_data(doc, headers=headers, rows=rows, font_size=8, header_color='2E75B6')

    # ===== Table S3 =====
    add_heading_styled(doc,
        'Table S3. Baseline Characteristics (N=5,728) by DSI Quartile, Including SOFA',
        level=2, font_size=11)
    headers, rows = read_csv_data('Table_S3_Baseline_Characteristics.csv')
    add_table_from_data(doc, headers=headers, rows=rows, font_size=8)

    # ===== Table S4 =====
    add_heading_styled(doc, 'Table S4. TRIPOD+AI Checklist (27 Items)', level=2, font_size=11)
    headers, rows = read_csv_data('Table_S4_TRIPOD_AI_Checklist.csv')
    add_table_from_data(doc, headers=headers, rows=rows, font_size=8, header_color='2E75B6')

    # ===== Table S5 =====
    add_heading_styled(doc,
        'Table S5. ICD Code Composition of \u201cOther\u201d Subtype (N=1,712)',
        level=2, font_size=11)
    headers, rows = read_csv_data('Table_S5_Other_Subtype_ICD.csv')
    add_table_from_data(doc, headers=headers, rows=rows[:15], font_size=8,
                        caption='Top 15 most frequent ICD-10 codes in the \u201cOther\u201d subtype')
    add_styled_paragraph(doc,
        '(Full list contains 30 codes; see Table_S5_Other_Subtype_ICD.csv for complete data.)',
        italic=True, font_size=9)

    # ===== Table S6 =====
    add_heading_styled(doc,
        'Table S6. Full Model Coefficients (Primary Model Without Surgery + DSI) Per TRIPOD+AI Guidelines',
        level=2, font_size=11)

    primary = v6_stats['P0_4_surgery_variations']['extended_no_surgery_DSI']
    coeffs = primary['coefficients']

    s6_rows = []
    s6_rows.append([
        'Intercept',
        str(round(primary['intercept'], 4)),
        str(round(primary['intercept_se'], 4)),
        '\u2014', '\u2014', '\u2014', '\u2014'
    ])
    for var_key, var_data in coeffs.items():
        var_label = var_names.get(var_key, var_key)
        s6_rows.append([
            var_label,
            str(round(var_data['beta'], 4)),
            str(round(var_data['se'], 4)),
            str(round(var_data['OR'], 3)),
            f"{round(var_data['CI_lower'], 3)}\u2013{round(var_data['CI_upper'], 3)}",
            var_data['P'],
            ''
        ])

    # Fix: merge CI into one column display
    s6_rows_fixed = []
    for row in s6_rows:
        s6_rows_fixed.append([row[0], row[1], row[2], row[3], row[4], row[5]])

    add_table_from_data(doc,
        headers=['Variable', '\u03b2', 'SE', 'OR', '95% CI', 'P'],
        rows=s6_rows_fixed,
        font_size=9,
        caption=f"Model: Extended baseline (without surgery) + DSI. AUC=0.790, Brier=0.127, "
                f"N=5,728, events=1,141")

    # ===== Table S7 =====
    add_heading_styled(doc,
        'Table S7. eICU-CRD Baseline Characteristics (N=5,755) by DSI Quartile',
        level=2, font_size=11)
    headers, rows = read_csv_data('Table_S7_eICU_Baseline_by_DSI_Quartile.csv')
    add_table_from_data(doc, headers=headers, rows=rows, font_size=8)

    # ===== Table S8 =====
    add_heading_styled(doc,
        'Table S8. Comparison of Excluded (N=3,205) vs Complete-Case (N=5,728) Patient Characteristics, '
        'Demonstrating Selection Bias',
        level=2, font_size=11)
    headers, rows = read_csv_data('Table_S8_Excluded_vs_CC_Characteristics.csv')
    add_table_from_data(doc, headers=headers, rows=rows, font_size=9)

    add_styled_paragraph(doc,
        'Note: The 36% exclusion rate was primarily driven by lactate non-availability (99% of excluded '
        'patients lacked lactate values). Excluded patients had substantially lower severity, confirming '
        'selection bias toward more severely ill patients who received arterial blood gas monitoring. '
        'Multiple imputation on N=8,933 addressed this bias.',
        italic=True, font_size=9, space_before=Pt(6))

    # ===== Table S9 (NEW) =====
    add_heading_styled(doc,
        'Table S9. Alternative Model (With Surgery) Coefficients and Performance',
        level=2, font_size=11)

    alt_model = v6_stats['P0_4_surgery_variations']['extended_with_surgery_DSI']
    alt_baseline = v6_stats['P0_4_surgery_variations']['extended_with_surgery_baseline']
    alt_coeffs = alt_model['coefficients']

    s9_rows = []
    s9_rows.append([
        'Intercept',
        str(round(alt_model['intercept'], 4)),
        str(round(alt_model['intercept_se'], 4)),
        '\u2014', '\u2014', '\u2014'
    ])
    for var_key, var_data in alt_coeffs.items():
        var_label = var_names.get(var_key, var_key)
        s9_rows.append([
            var_label,
            str(round(var_data['beta'], 4)),
            str(round(var_data['OR'], 3)),
            f"{round(var_data['CI_lower'], 3)}\u2013{round(var_data['CI_upper'], 3)}",
            var_data['P'],
            ''
        ])

    add_table_from_data(doc,
        headers=['Variable', '\u03b2', 'OR', '95% CI', 'P', ''],
        rows=s9_rows,
        font_size=9,
        caption=f"Alternative model: Extended baseline (with surgery) + DSI. "
                f"Baseline AUC={alt_baseline['AUC']:.3f}, DSI AUC={alt_model['AUC']:.3f}, "
                f"\u0394AUC={v6_stats['delta_auc_summary']['with_surgery']:.3f}, "
                f"Brier={alt_model['Brier']:.3f}")

    add_styled_paragraph(doc,
        'Note: Surgery appeared protective (OR=0.68) but likely reflects survivorship bias: only 5.1% '
        'had surgery \u226424h from ICU admission vs 67.4% \u201cduring hospitalization.\u201d Including surgery '
        'increased baseline AUC by only 0.002. The primary model (without surgery) is preferred.',
        italic=True, font_size=9, space_before=Pt(6))

    # ===== Table S10 (NEW) =====
    add_heading_styled(doc,
        'Table S10. Parsimonious Model (Without Vasopressor/MV) Performance',
        level=2, font_size=11)

    pars_baseline = v6_stats['P1_4_parsimonious']['parsimonious_baseline']
    pars_dsi = v6_stats['P1_4_parsimonious']['parsimonious_DSI']
    pars_coeffs = pars_dsi['coefficients']

    s10_rows = []
    s10_rows.append([
        'Intercept',
        str(round(pars_dsi['intercept'], 4)),
        str(round(pars_dsi['intercept_se'], 4)),
        '\u2014', '\u2014', '\u2014'
    ])
    for var_key, var_data in pars_coeffs.items():
        var_label = var_names.get(var_key, var_key)
        s10_rows.append([
            var_label,
            str(round(var_data['beta'], 4)),
            str(round(var_data['OR'], 3)),
            f"{round(var_data['CI_lower'], 3)}\u2013{round(var_data['CI_upper'], 3)}",
            var_data['P'],
            ''
        ])

    add_table_from_data(doc,
        headers=['Variable', '\u03b2', 'OR', '95% CI', 'P', ''],
        rows=s10_rows,
        font_size=9,
        caption=f"Parsimonious model: age + sex + CCI + lactate + WBC + SOFA + DSI. "
                f"Baseline AUC={pars_baseline['AUC']:.3f}, DSI AUC={pars_dsi['AUC']:.3f}, "
                f"\u0394AUC={v6_stats['delta_auc_summary']['parsimonious']:.3f}, "
                f"Brier={pars_dsi['Brier']:.3f}")

    add_styled_paragraph(doc,
        'Note: Vasopressor use (P=0.14) and mechanical ventilation (P=0.45) were non-significant after '
        'SOFA adjustment in the primary model, consistent with SOFA absorbing their predictive '
        'information. The parsimonious model confirms that removing these redundant covariates does not '
        'affect DSI\'s independent predictive value.',
        italic=True, font_size=9, space_before=Pt(6))

    # ===== Table S11 (NEW) =====
    add_heading_styled(doc,
        'Table S11. Sensitivity Analyses Summary (12 Scenarios)',
        level=2, font_size=11)

    # Build summary from v6_stats
    sv = v6_stats['P0_4_surgery_variations']
    pe = v6_stats['P0_5_subtype_exclusions']
    pa = v6_stats['P1_4_parsimonious']

    s11_rows = [
        ['1', 'Primary model (no surgery)', '5,728',
         f"{sv['extended_no_surgery_DSI']['AUC']:.3f}",
         f"{sv['extended_no_surgery_DSI']['coefficients']['DSI_mean']['OR']:.2f}",
         sv['extended_no_surgery_DSI']['coefficients']['DSI_mean']['P'],
         'Primary model'],
        ['2', 'Model with surgery', '5,728',
         f"{sv['extended_with_surgery_DSI']['AUC']:.3f}",
         f"{sv['extended_with_surgery_DSI']['coefficients']['DSI_mean']['OR']:.2f}",
         sv['extended_with_surgery_DSI']['coefficients']['DSI_mean']['P'],
         'Surgery OR=0.68 (survivorship bias)'],
        ['3', 'Surgery \u226424h model', '5,728',
         f"{sv['extended_surgery24h_DSI']['AUC']:.3f}",
         f"{sv['extended_surgery24h_DSI']['coefficients']['DSI_mean']['OR']:.2f}",
         sv['extended_surgery24h_DSI']['coefficients']['DSI_mean']['P'],
         'surgery_24h OR=0.88, P=0.46'],
        ['4', 'Parsimonious (no vasopressor/MV)', '5,728',
         f"{pa['parsimonious_DSI']['AUC']:.3f}",
         f"{pa['parsimonious_DSI']['coefficients']['DSI_mean']['OR']:.2f}",
         pa['parsimonious_DSI']['coefficients']['DSI_mean']['P'],
         'Redundant covariates removed'],
        ['5', 'No vasopressor/MV + surgery', '5,728',
         f"{pa['extended_no_VMV_baseline']['AUC']:.3f}",
         f"{pa['extended_no_VMV_baseline']['coefficients']['DSI_mean']['OR']:.2f}",
         pa['extended_no_VMV_baseline']['coefficients']['DSI_mean']['P'],
         'Alternative parsimonious'],
        ['6', 'Excluding \u201cOther\u201d subtype (with surgery)', '4,016',
         f"{pe['excl_other_DSI']['AUC']:.3f}",
         f"{pe['excl_other_DSI']['coefficients']['DSI_mean']['OR']:.2f}",
         pe['excl_other_DSI']['coefficients']['DSI_mean']['P'],
         'Primary subtypes only'],
        ['7', 'Excluding \u201cOther\u201d (no surgery)', '4,016',
         f"{pe['excl_other_no_surgery_DSI']['AUC']:.3f}",
         f"{pe['excl_other_no_surgery_DSI']['coefficients']['DSI_mean']['OR']:.2f}",
         pe['excl_other_no_surgery_DSI']['coefficients']['DSI_mean']['P'],
         'Primary subtypes only'],
        ['8', 'Non-surgical subgroup', '1,865', '0.826', '2.28', '<10\u207b\u00b9\u00b9',
         'Best performance'],
        ['9', 'Surgical subgroup', '3,863', '0.777', '2.22', '<10\u207b\u00b9\u00b9', ''],
        ['10', 'Multiple imputation', '8,933', '0.822', '2.65', '<10\u207b\u00b9\u00b9',
         'Addresses selection bias'],
        ['11', 'MICU/SICU/TSICU only', '3,594', '0.800', '2.14', '<10\u207b\u00b9\u00b9', ''],
        ['12', 'eICU-CRD external validation', '5,755', '0.792', '\u2014', '0.0026 (DeLong)',
         '208 hospitals; recalibration required'],
    ]

    add_table_from_data(doc,
        headers=['#', 'Analysis', 'N', 'AUC', 'DSI OR', 'P', 'Notes'],
        rows=s11_rows,
        font_size=8,
        caption='DSI\'s independent predictive value was robust across all 12 sensitivity analyses '
                '(OR range 2.15\u20132.65, all P<10\u207b\u00b9\u00b9)')

    # ===== Figure Legends S1-S7 =====
    add_heading_styled(doc, 'Figure Legends', level=1, font_size=14)

    fig_s_items = [
        ('Figure S1. Calibration Plots for Basic Baseline Models',
         'Calibration plots for basic baseline (age+sex+CCI), showing poor calibration '
         '(HL P=0.016, Brier=0.155).'),
        ('Figure S2. Kaplan-Meier Curves by DSI Quartile',
         'Kaplan-Meier survival curves stratified by DSI quartile, using hospital length of stay '
         'as the time axis. Log-rank \u03c7\u00b2=71.2, P=2.33\u00d710\u207b\u00b9\u2075. These KM curves are a visual '
         'supplement; the primary competing risk analysis (CIF, Figure S7) is more rigorous.'),
        ('Figure S3. RCS Dose-Response Curves for SI, MSI, DSI, Age-SI',
         'Restricted cubic spline (4 knots) dose-response curves for all four SI derivatives. '
         'All P_overall<0.001; all P_nonlinear>0.05, supporting linear dose-response relationships.'),
        ('Figure S4. DCA Net Benefit Across Threshold Probabilities',
         'Decision curve analysis showing net benefit across threshold probabilities (1\u201350%). '
         'The extended+DSI model provided modest incremental net benefit over the extended baseline '
         'at clinically relevant thresholds (5\u201325%).'),
        ('Figure S5. Calibration Plots (4 Model Levels)',
         'Calibration plots for basic baseline, extended baseline (no surgery), extended+DSI (no '
         'surgery), and extended+all SI derivatives. Extended+DSI: HL P=0.691, Brier=0.126.'),
        ('Figure S6. Subgroup ROC Curves by Subtype',
         'ROC curves for subtype-specific models: inflammation (AUC=0.819), obstruction (AUC=0.749), '
         'perforation (AUC=0.766), ischemia (AUC=0.807). Ischemia showed the highest mortality (40.5%).'),
        ('Figure S7. Cumulative Incidence Functions by DSI Quartile',
         'Cumulative incidence functions for in-hospital death (competing risk: discharge alive) by '
         'DSI quartile, demonstrating progressive risk across quartiles. CIF curves are descriptive; '
         'formal Fine-Gray subdistribution hazard models were not implemented.'),
    ]
    for title, desc in fig_s_items:
        add_styled_paragraph(doc, title, bold=True, font_size=12, space_after=Pt(2))
        add_styled_paragraph(doc, desc, font_size=12, space_after=Pt(6))

    # Save
    out_path = os.path.join(BASE_DIR, 'Supplementary_Materials_AIC_v7.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


# ============================================================
# 3. Generate Cover_Letter_AIC_v7.docx
# ============================================================

def generate_cover_letter():
    doc = Document()
    setup_document(doc)

    # Date
    add_styled_paragraph(doc, 'July 20, 2026', font_size=12, space_after=Pt(12))

    # Recipient
    add_styled_paragraph(doc,
        'The Editor-in-Chief\nAnnals of Intensive Care',
        font_size=12, space_after=Pt(12))

    # Re line
    add_styled_paragraph(doc,
        'Re: Submission of original research article \u2014 \u201cDiastolic Shock Index as an Independent '
        'Predictor of In-Hospital Mortality in Critically Ill Patients with Acute Abdomen: A '
        'Retrospective Cohort Study with External Validation\u201d',
        bold=True, font_size=12, space_after=Pt(12))

    # Salutation
    add_styled_paragraph(doc, 'Dear Editor,', font_size=12, space_after=Pt(6))

    # Opening
    add_styled_paragraph(doc,
        'We are pleased to submit our original research article entitled \u201cDiastolic Shock Index as '
        'an Independent Predictor of In-Hospital Mortality in Critically Ill Patients with Acute '
        'Abdomen: A Retrospective Cohort Study with External Validation\u201d for consideration for '
        'publication in Annals of Intensive Care.',
        font_size=12, first_line_indent=Inches(0.5))

    # Summary
    add_heading_styled(doc, 'Summary of the Study', level=2, font_size=12)
    add_styled_paragraph(doc,
        'This retrospective cohort study utilized the MIMIC-IV v3.1 database (546,028 total admissions, '
        '2008\u20132022) to evaluate the prognostic value of four shock index-derived parameters for '
        'predicting in-hospital mortality in critically ill patients with acute abdomen. Among 5,728 '
        'complete-case ICU stays (median age 68 years, 56.0% male, in-hospital mortality 19.9%), '
        'diastolic shock index (DSI = HR/DBP, 24h mean) emerged as the strongest SI derivative, '
        'maintaining independent prognostic value (OR=2.18, 95% CI 1.79\u20132.65, P=7.59\u00d710\u207b\u00b9\u2075) '
        'even after adjustment for the SOFA score and established ICU covariates. The primary extended '
        'baseline model excluded abdominal surgery to avoid survivorship bias (67.4% had surgery '
        '\u201cduring hospitalization\u201d but only 5.1% \u226424h). Adding DSI yielded AUC=0.790 vs baseline '
        '0.785 (\u0394AUC=0.005, DeLong P=0.012).',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc,
        'Crucially, we transparently acknowledge that \u0394AUC=0.005 is below conventional clinical '
        'relevance thresholds (\u22650.02 per Cook and Vickers), and that the categorical NRI crosses zero '
        '(0.008, 95% CI \u22120.009 to 0.044). We therefore position DSI not as a SOFA replacement, but as '
        'a complementary zero-cost bedside tool. Its clinical value lies in: (1) the dramatic DSI '
        'quartile mortality gradient (12.1%\u219232.8%, 2.7-fold, P=2.02\u00d710\u207b\u2074\u2079), providing '
        'immediately actionable risk thresholds from routinely monitored HR and DBP\u2014available without '
        'laboratory turnaround time; (2) the significant category-free NRI (0.252) and IDI (0.013), '
        'confirming additional continuous prognostic information; and (3) the independent OR persisting '
        'after SOFA adjustment. DSI provides immediate risk stratification when SOFA components '
        '(platelets, bilirubin, PaO\u2082, vasopressor doses) are pending.',
        font_size=12, first_line_indent=Inches(0.5))

    # External validation
    add_heading_styled(doc, 'External Validation', level=2, font_size=12)
    add_styled_paragraph(doc,
        'External validation was performed in eICU-CRD v2.0 (N=5,755, 208 US hospitals, 2014\u20132015). '
        'Discrimination was preserved (AUC=0.792, \u0394AUC=0.0074, DeLong P=0.0026), and the DSI quartile '
        'gradient was closely replicated (12.0%\u219233.5%). However, direct application yielded '
        'catastrophically poor calibration (Brier 0.38\u20130.59), requiring logistic recalibration '
        '(intercept shift \u22123.935, slope 0.952). The near-ideal slope confirms discrimination '
        'transportability, but the large intercept shift means absolute risk predictions require local '
        'recalibration\u2014a finding we report transparently per TRIPOD+AI guidelines.',
        font_size=12, first_line_indent=Inches(0.5))

    # Why AIC
    add_heading_styled(doc, 'Why Annals of Intensive Care', level=2, font_size=12)
    add_styled_paragraph(doc,
        'Annals of Intensive Care is the ideal venue for this work because DSI was originally described '
        'in septic shock by Ospina-Tasc\u00f3n et al. (Ann Intensive Care 2020;10:41), making our extension '
        'to acute abdomen a natural continuation of work published in your journal. The study addresses '
        'ICU risk stratification\u2014a core focus of Annals of Intensive Care\u2014and provides rigorous '
        'external validation in a multi-center database (208 hospitals). The transparent reporting of '
        'both strengths (independent OR, dramatic quartile gradient, external validation preserved) and '
        'limitations (\u0394AUC below clinical thresholds, categorical NRI crossing zero, recalibration '
        'required) aligns with the journal\'s emphasis on honest, balanced reporting.',
        font_size=12, first_line_indent=Inches(0.5))

    # Relationship disclosure
    add_heading_styled(doc, 'Disclosure Regarding Editor-in-Chief', level=2, font_size=12)
    add_styled_paragraph(doc,
        'We wish to transparently declare that Jean-Louis Teboul, the Editor-in-Chief of Annals of '
        'Intensive Care, is a co-author of the original DSI paper (Ospina-Tasc\u00f3n GA, Teboul JL, '
        'Hernandez G, et al. Diastolic shock index and clinical outcomes in septic shock. Ann Intensive '
        'Care. 2020;10:41), which is cited as reference 5 in our manuscript. The authors of the present '
        'study have no personal or professional relationship with Professor Teboul beyond this '
        'bibliographic connection. We declare this relationship in the interest of full transparency '
        'and request that the editorial handling process proceed according to the journal\'s standard '
        'conflict-of-interest policies.',
        font_size=12, first_line_indent=Inches(0.5))

    # AI use
    add_heading_styled(doc, 'Use of AI', level=2, font_size=12)
    add_styled_paragraph(doc,
        'During the preparation of this manuscript, the authors used a large language model for language '
        'polishing and manuscript editing only. No AI tools were used for data extraction, statistical '
        'analysis, or scientific interpretation. All data extraction, statistical analyses, figure '
        'generation, and scientific interpretation were performed independently by the authors. The '
        'authors reviewed and edited all AI-assisted content and take full responsibility for the '
        'content of the publication.',
        font_size=12, first_line_indent=Inches(0.5))

    # Closing
    add_heading_styled(doc, 'Closing', level=2, font_size=12)
    add_styled_paragraph(doc,
        'This manuscript has not been published and is not under consideration elsewhere. Both authors '
        'approve the manuscript and have no conflicts of interest. We believe this study makes a '
        'meaningful contribution to ICU risk stratification in acute abdomen by establishing DSI as an '
        'independent, zero-cost bedside predictor with externally validated discrimination and a '
        'dramatic quartile mortality gradient.',
        font_size=12, first_line_indent=Inches(0.5))

    add_styled_paragraph(doc, 'Sincerely,', font_size=12, space_before=Pt(12))
    add_styled_paragraph(doc, 'Dengke Wu, MD', font_size=12, bold=True)
    add_styled_paragraph(doc,
        'Department of Emergency Medicine\nThe Second Xiangya Hospital of Central South University\n'
        'Changsha 410011, Hunan, China\nwudk2010@csu.edu.cn',
        font_size=12)

    # Save
    out_path = os.path.join(BASE_DIR, 'Cover_Letter_AIC_v7.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path


# ============================================================
# Main
# ============================================================

if __name__ == '__main__':
    print('Generating SCI_paper_v7.docx...')
    p1 = generate_main_paper()

    print('Generating Supplementary_Materials_AIC_v7.docx...')
    p2 = generate_supplementary()

    print('Generating Cover_Letter_AIC_v7.docx...')
    p3 = generate_cover_letter()

    print(f'\nAll 3 DOCX files generated:')
    print(f'  1. {p1}')
    print(f'  2. {p2}')
    print(f'  3. {p3}')
