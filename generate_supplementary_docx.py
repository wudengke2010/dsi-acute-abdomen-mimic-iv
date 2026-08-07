"""Generate combined Supplementary Materials Word document for AIC submission."""
import csv
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# Title page
title = doc.add_heading('Supplementary Materials', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    'Shock Index-Derived Parameters as Predictors of In-Hospital Mortality '
    'in Critically Ill Patients with Acute Abdomen: '
    'A Retrospective Cohort Study from MIMIC-IV'
)
run.italic = True
run.font.size = Pt(12)

doc.add_paragraph()  # spacer

# Helper function to add a CSV as a table
def add_csv_table(doc, csv_path, table_title, max_col_width=None):
    doc.add_page_break()
    doc.add_heading(table_title, level=2)

    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        doc.add_paragraph('(Empty)')
        return

    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            if i == 0:  # header row
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacer


# Helper to add figure
def add_figure(doc, img_path, fig_title, caption, width_inches=6.0):
    doc.add_page_break()
    doc.add_heading(fig_title, level=2)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f'(Image not found: {img_path})')
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)


BASE = os.path.dirname(os.path.abspath(__file__))

# Table S1
add_csv_table(doc, os.path.join(BASE, 'Table_S1_ICD_Codes.csv'),
              'Table S1. ICD-9 and ICD-10 Codes for Acute Abdomen Identification and Subtype Classification')

# Table S2
add_csv_table(doc, os.path.join(BASE, 'Table_S2_STROBE_Checklist.csv'),
              'Table S2. STROBE Checklist for Reporting Observational Studies')

# Table S3
add_csv_table(doc, os.path.join(BASE, 'Table_S3_Baseline_Characteristics.csv'),
              'Table S3. Baseline Characteristics of the Complete-Case Cohort (N=5,728) by DSI Quartile')

# Table S4
add_csv_table(doc, os.path.join(BASE, 'Table_S4_TRIPOD_AI_Checklist.csv'),
              'Table S4. TRIPOD+AI Checklist (27 Items) for Reporting Clinical Prediction Models')

# Table S5
add_csv_table(doc, os.path.join(BASE, 'Table_S5_Other_Subtype_ICD.csv'),
              'Table S5. ICD Code Composition of the "Other" Acute Abdomen Subtype (N=1,712)')

# Table S6
add_csv_table(doc, os.path.join(BASE, 'Table_S6_Model_Coefficients.csv'),
              'Table S6. Full Model Coefficients for the Extended Baseline + DSI Model')

# Figure S1
add_figure(doc, os.path.join(BASE, 'figures', 'FigS1_Calibration_basic.png'),
           'Figure S1',
           'Calibration plots for basic baseline, basic baseline + DSI, and basic baseline + all SI derivatives models.')

# Figure S2
add_figure(doc, os.path.join(BASE, 'figures', 'FigS2_KM.png'),
           'Figure S2',
           'Kaplan-Meier in-hospital survival curves by DSI quartile (hospital length of stay as time axis). '
           'Log-rank P=2.33×10⁻¹⁵. The primary competing risk analysis is presented in Figure 8.')

output_path = os.path.join(BASE, 'Supplementary_Materials_AIC.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
print(f'Size: {os.path.getsize(output_path):,} bytes')
