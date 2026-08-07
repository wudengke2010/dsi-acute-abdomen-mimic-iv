"""
Convert SCI_paper_v8.md to DOCX with AIC formatting:
- Double-line spacing
- Continuous line numbers
- Times New Roman 12pt (body) / 10pt (tables/references)
- 1-inch margins, A4
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))

def add_line_numbers(section):
    """Add continuous line numbering to a section via XML."""
    sectPr = section._sectPr
    # Remove existing lnNumType
    for ln in sectPr.findall(qn('w:lnNumType')):
        sectPr.remove(ln)
    lnNumType = OxmlElement('w:lnNumType')
    lnNumType.set(qn('w:countBy'), '1')
    lnNumType.set(qn('w:restart'), 'continuous')
    # Insert after pgMar
    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.addnext(lnNumType)
    else:
        sectPr.append(lnNumType)

def format_document(doc):
    """Apply AIC formatting to a Document."""
    # Set page size to A4
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        add_line_numbers(section)

    # Set Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)

    # Apply to all paragraphs
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        for run in p.runs:
            run.font.name = 'Times New Roman'
            if run.font.size is None or run.font.size == Pt(12):
                run.font.size = Pt(12)

    # Apply to all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        if run.font.size is None or run.font.size >= Pt(11):
                            run.font.size = Pt(10)

def add_table_from_md(doc, header_line, rows):
    headers = [c.strip() for c in header_line.split('|')[1:-1]]
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    for r_idx, row_line in enumerate(rows):
        cells_data = [c.strip() for c in row_line.split('|')[1:-1]]
        for c_idx, cell_text in enumerate(cells_data):
            if c_idx < cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
    doc.add_paragraph()

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_header = None
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        if not line.strip():
            if in_table and table_header and table_rows:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False
                table_header = None
                table_rows = []
            i += 1
            continue

        if line.startswith('#'):
            if in_table and table_header and table_rows:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False
                table_header = None
                table_rows = []
            level = line.count('#', 0, line.find(' '))
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        if '|' in line and line.strip().startswith('|'):
            if re.match(r'^\|[\s\-:]+\|', line):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = line
                table_rows = []
            else:
                table_rows.append(line)
            i += 1
            continue
        else:
            if in_table and table_header and table_rows:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False
                table_header = None
                table_rows = []

        if line.strip() == '---':
            i += 1
            continue

        # Bold paragraphs and regular text
        if line.startswith('**') and '**' in line[2:]:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
            i += 1
            continue

        if line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            num = int(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            run = p.add_run(f"{num}. {text}")
            run.font.size = Pt(10)
            i += 1
            continue

        p = doc.add_paragraph(line)
        i += 1

    if in_table and table_header and table_rows:
        add_table_from_md(doc, table_header, table_rows)

    # Apply AIC formatting
    format_document(doc)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

# Generate main manuscript
convert_md_to_docx(
    os.path.join(BASE, 'SCI_paper_v8.md'),
    os.path.join(BASE, 'SCI_paper_v8.docx')
)

# Generate supplementary materials DOCX
supp_md = os.path.join(BASE, 'Supplementary_Materials_v8.md')
if os.path.exists(supp_md):
    convert_md_to_docx(supp_md, os.path.join(BASE, 'Supplementary_Materials_v8.docx'))
else:
    # Create supplementary from v7 supplementary content
    print("Supplementary_Materials_v8.md not found, creating from v7...")
    # Read v8 paper to extract supplementary section
    with open(os.path.join(BASE, 'SCI_paper_v8.md'), 'r', encoding='utf-8') as f:
        v8_text = f.read()

    # Extract supplementary section
    supp_match = re.search(r'## Supplementary Materials\n\n(.*?)(?=\n---\n\n## Figure Legends)', v8_text, re.DOTALL)
    supp_content = supp_match.group(1) if supp_match else ''

    # Create supplementary document
    supp_doc = Document()
    style = supp_doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    supp_doc.add_heading('Supplementary Materials', level=1)
    supp_doc.add_heading('Diastolic Shock Index as an Independent Predictor of In-Hospital Mortality in Critically Ill Patients with Acute Abdomen', level=2)

    # Parse supplementary items
    for line in supp_content.strip().split('\n'):
        if line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            p = supp_doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        elif line.strip():
            supp_doc.add_paragraph(line.strip())

    # Add Table S9, S10, S11 from CSV files
    import csv
    for table_name, csv_file in [
        ('Table S9. Alternative Model (with Surgery) Coefficients and Performance', 'Table_S9_Alternative_Model_with_Surgery.csv'),
        ('Table S10. Parsimonious Model (without Vasopressor/MV) Performance', 'Table_S10_Parsimonious_Model.csv'),
        ('Table S11. Sensitivity Analyses Summary (14 Scenarios)', 'Table_S11_Sensitivity_Analyses_Summary.csv'),
        ('Table S12. DSI Threshold Diagnostic Performance', 'DSI_diagnostic_performance.csv'),
        ('Table S13. E-value Analysis for DSI Association with In-Hospital Mortality', 'Table_S13_Evalue.csv'),
    ]:
        csv_path = os.path.join(BASE, csv_file)
        if os.path.exists(csv_path):
            supp_doc.add_heading(table_name, level=3)
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
                if rows:
                    table = supp_doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell_text in enumerate(row):
                            if c_idx < len(table.rows[r_idx].cells):
                                cell = table.rows[r_idx].cells[c_idx]
                                cell.text = cell_text
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.size = Pt(9)
                                        run.font.name = 'Times New Roman'
                                        if r_idx == 0:
                                            run.bold = True
            supp_doc.add_paragraph()

    # Add component decomposition results
    supp_doc.add_heading('Component Decomposition Analysis Results', level=3)
    import json
    comp_path = os.path.join(BASE, 'component_decomposition_results.json')
    if os.path.exists(comp_path):
        with open(comp_path, 'r', encoding='utf-8') as f:
            comp = json.load(f)
        for key, val in comp.items():
            p = supp_doc.add_paragraph(f"{key}: {val}")
            for run in p.runs:
                run.font.size = Pt(10)

    format_document(supp_doc)
    supp_doc.save(os.path.join(BASE, 'Supplementary_Materials_v8.docx'))
    print(f"Saved: {os.path.join(BASE, 'Supplementary_Materials_v8.docx')}")

# Generate Cover Letter v8
cover_content = """Date: August 7, 2026

To: The Editor-in-Chief, Annals of Intensive Care

Re: Submission of original research article — "Diastolic Shock Index as an Independent Predictor of In-Hospital Mortality in Critically Ill Patients with Acute Abdomen: A Retrospective Cohort Study with External Validation"

Dear Editor,

We are pleased to submit our original research article entitled "Diastolic Shock Index as an Independent Predictor of In-Hospital Mortality in Critically Ill Patients with Acute Abdomen: A Retrospective Cohort Study with External Validation" for consideration for publication in Annals of Intensive Care.

This retrospective cohort study utilized MIMIC-IV v3.1 (546,028 admissions, 2008-2022) to evaluate DSI (HR/DBP) as an independent, zero-cost bedside predictor of in-hospital mortality in critically ill patients with acute abdomen, with external validation in eICU-CRD. Among 5,728 complete-case ICU stays (in-hospital mortality 19.9%), DSI maintained independent prognostic value (OR=2.18, 95% CI 1.79-2.65, P=7.59x10^-15) after SOFA adjustment. The primary model excluded surgery to avoid survivorship bias (only 5.1% had surgery <=24h). Adding DSI yielded AUC=0.790 vs baseline 0.785 (DeltaAUC=0.005, DeLong P=0.012). We transparently acknowledge that DeltaAUC is below clinical relevance thresholds (>=0.02 per Cook 2007, Vickers et al. 2011), and categorical NRI crosses zero (0.008, 95% CI -0.009 to 0.044). DSI's clinical value lies in its independent association (OR=2.18), dramatic quartile mortality gradient (12.1%->32.8%, 2.7-fold), and zero-cost bedside availability complementary to SOFA when laboratory data are pending. Component decomposition confirmed that DSI captures prognostic information unavailable from isolated HR (AUC=0.571) or DBP (AUC=0.597).

Robustness was confirmed via an a priori DAG causal framework, E-value analysis (point estimate 3.78; 2.98 for CI lower bound, indicating an unmeasured confounder would need RR >3.8 to explain the association), Benjamini-Hochberg FDR correction for multiplicity, multiple imputation (N=8,933, AUC=0.822), MICE comparison (DSI OR consistent: IterativeImputer 2.65 vs MICE 2.63), and 14 sensitivity analyses addressing nine bias categories. External validation across 208 hospitals preserved discrimination (AUC=0.792) and replicated the quartile gradient (12.0%->33.5%), though calibration required local recalibration (intercept shift -3.935; slope 0.952). Analysis code is publicly available at https://github.com/wudengke2010/dsi-acute-abdomen-mimic-iv.

Why This Study Fits Annals of Intensive Care:

1. Novelty: First systematic evaluation of DSI — originally described by Ospina-Tascon et al. in Annals of Intensive Care (2020) for septic shock — in acute abdomen ICU patients. Component decomposition extends the principle that DSI outperforms isolated HR/DBP to a new population.

2. Clinical relevance: Acute abdomen requiring ICU admission carries 15-20% mortality. A zero-cost bedside tool providing immediate risk stratification before laboratory results benefits intensivists directly.

3. Methodological rigor: STROBE and TRIPOD+AI-compliant; DAG-based causal framework; transparent disclosure of DeltaAUC below clinical thresholds; Benjamini-Hochberg FDR correction; two multiple imputation strategies; 14 sensitivity analyses across nine bias categories; external validation with both un-recalibrated and recalibrated metrics.

We note that Professor Jean-Louis Teboul, Editor-in-Chief of Annals of Intensive Care, is a co-author of the original DSI paper (Ospina-Tascon et al. 2020). We believe our work represents a meaningful extension of that foundational research to a previously unexamined population.

This manuscript has not been published previously and is not under consideration elsewhere. All authors have approved the final manuscript. The study uses publicly available, de-identified data. The authors declare no conflicts of interest. This work was supported by the National Health Commission (GWJJMB202510024181), Changsha Science and Technology Bureau (kq2014242), and Hunan Provincial Natural Science Foundation (2021JJ30959). Funders had no role in study design, analysis, or publication.

Sincerely,

Dengke Wu, MD
Department of Emergency Medicine, and Emergency Medicine and Difficult Diseases Institute
The Second Xiangya Hospital of Central South University
Changsha 410011, Hunan, China
Email: wudk2010@csu.edu.cn
ORCID: 0009-0008-1363-9621

Jiqiang Liu
Department of Emergency Medicine
The Second Xiangya Hospital of Central South University
Changsha 410011, Hunan, China
ORCID: 0009-0000-9884-3089
"""

cover_doc = Document()
style = cover_doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for line in cover_content.strip().split('\n'):
    if line.strip():
        cover_doc.add_paragraph(line)
    else:
        cover_doc.add_paragraph()

format_document(cover_doc)
cover_doc.save(os.path.join(BASE, 'Cover_Letter_v8.docx'))
print(f"Saved: {os.path.join(BASE, 'Cover_Letter_v8.docx')}")

print("\n=== All v8 DOCX files generated ===")
