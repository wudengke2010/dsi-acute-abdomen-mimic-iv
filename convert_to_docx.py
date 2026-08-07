"""Convert SCI_paper_v5_revised.md to Word .docx for Elsevier submission."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_heading_from_md(doc, line):
    """Parse markdown heading and add to document."""
    level = line.count('#', 0, line.find(' '))
    text = line.lstrip('#').strip()
    h = doc.add_heading(text, level=min(level, 4))
    return h

def add_table_from_md(doc, header_line, rows):
    """Parse markdown table and add as Word table."""
    # Parse header
    headers = [c.strip() for c in header_line.split('|')[1:-1]]
    cols = len(headers)
    
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    # Data rows
    for r_idx, row_line in enumerate(rows):
        cells_data = [c.strip() for c in row_line.split('|')[1:-1]]
        for c_idx, cell_text in enumerate(cells_data):
            if c_idx < cols:
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    
    doc.add_paragraph()  # spacing after table

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_header = None
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Skip empty lines
        if not line.strip():
            if in_table and table_header and table_rows:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False
                table_header = None
                table_rows = []
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            if in_table and table_header and table_rows:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False
                table_header = None
                table_rows = []
            add_heading_from_md(doc, line)
            i += 1
            continue
        
        # Table detection
        if '|' in line and line.strip().startswith('|'):
            # Check if separator line
            if re.match(r'^\|[\s\-:]+\|', line):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = line
                table_rows = []
            else:
                table_rows.append(line)
            i += 1
            continue
        else:
            if in_table and table_header and table_rows:
                add_table_from_md(doc, table_header, table_rows)
                in_table = False
                table_header = None
                table_rows = []
        
        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('─' * 50)
            i += 1
            continue
        
        # Bold paragraphs
        if line.startswith('**') and '**' in line[2:]:
            p = doc.add_paragraph()
            # Parse bold segments
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
            i += 1
            continue
        
        # Bullet points
        if line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
            i += 1
            continue
        
        # References (numbered list)
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            num = int(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            run = p.add_run(f"{num}. {text}")
            run.font.size = Pt(10)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph(line)
        i += 1
    
    # Handle trailing table
    if in_table and table_header and table_rows:
        add_table_from_md(doc, table_header, table_rows)
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    convert_md_to_docx(
        os.path.join(base, 'SCI_paper_v5_revised.md'),
        os.path.join(base, 'SCI_paper_v5_revised.docx')
    )
