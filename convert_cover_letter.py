"""Convert Cover_Letter_AIC.md to Word .docx"""
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_cover_letter(md_path, docx_path):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.rstrip('\n')
        
        if not line.strip():
            doc.add_paragraph()
            continue
        
        if line.startswith('# '):
            text = line[2:]
            h = doc.add_heading(text, level=1)
        elif line.startswith('## '):
            text = line[3:]
            h = doc.add_heading(text, level=2)
        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('- '):
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. '):
            text = line[3:]
            p = doc.add_paragraph(style='List Number')
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        elif line.strip() == '---':
            p = doc.add_paragraph()
            run = p.add_run('─' * 60)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    convert_cover_letter(
        os.path.join(base, 'Cover_Letter_AIC.md'),
        os.path.join(base, 'Cover_Letter_AIC.docx')
    )
